#!/usr/bin/env python3
"""Generate the Kaft (cover page) from the Howest template.

Usage (from thesis/ directory):
    python3 generate_kaft.py

Requires python-docx (pip install python-docx).
"""

from docx import Document
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATE = SCRIPT_DIR / ".." / "template" / "Kaft_bachelorproef_2025_2026_EN.docx"
OUTPUT = SCRIPT_DIR / "build" / "Kaft_Snaet_2026.docx"

# --- Fill in your details here ---
TITLE = (
    "How Can a Semi-Supervised Neural Network Be Efficiently "
    "Implemented in Rust for the Automatic Labeling of Partially "
    "Labeled Datasets on an Edge Device?"
)
SUBTITLE = ""  # Optional, max 1 line
INTERNAL_PROMOTOR = "Gilles Depypere"
EXTERNAL_PROMOTOR = "Sandro Queirós"
STUDENT_NAME = "Warre Snaet"


def main():
    doc = Document(str(TEMPLATE))

    # Paragraph 0: Title (Kaft_Style1) — currently empty
    doc.paragraphs[0].text = TITLE

    # Paragraph 1: Subtitle (Kaft_Style1)
    if SUBTITLE:
        doc.paragraphs[1].text = SUBTITLE
    else:
        doc.paragraphs[1].text = ""

    # Paragraph 2: Promotors (Kaft_Style_Normal)
    promotor_text = f"internal promotor: {INTERNAL_PROMOTOR}"
    promotor_text += f"\nexternal promotor: {EXTERNAL_PROMOTOR}"
    doc.paragraphs[2].text = promotor_text

    # Paragraph 4: "Research question carried out by" — keep as is
    # Paragraph 5: Student name (Kaft_Style1) — currently empty
    doc.paragraphs[5].text = STUDENT_NAME

    # Paragraphs 6-8: fixed text, keep as is

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Kaft generated: {OUTPUT}")


if __name__ == "__main__":
    main()
